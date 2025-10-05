#!/usr/bin/env python3
"""
Create a Word document with resized screenshots from the CloudSecure AI project
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import glob

def resize_image(image_path, max_width=6, max_height=4):
    """Resize image to fit within specified dimensions while maintaining aspect ratio"""
    try:
        with Image.open(image_path) as img:
            # Calculate new dimensions maintaining aspect ratio
            width, height = img.size
            aspect_ratio = width / height
            
            if width > height:
                new_width = min(max_width, width)
                new_height = new_width / aspect_ratio
                if new_height > max_height:
                    new_height = max_height
                    new_width = new_height * aspect_ratio
            else:
                new_height = min(max_height, height)
                new_width = new_height * aspect_ratio
                if new_width > max_width:
                    new_width = max_width
                    new_height = new_width / aspect_ratio
            
            # Resize the image
            resized_img = img.resize((int(new_width), int(new_height)), Image.Resampling.LANCZOS)
            
            # Save resized image
            resized_path = f"resized_{os.path.basename(image_path)}"
            resized_img.save(resized_path)
            return resized_path
    except Exception as e:
        print(f"Error resizing {image_path}: {e}")
        return image_path

def create_word_document():
    """Create a Word document with screenshots and descriptions"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('CloudSecure AI - Screenshots Gallery', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run("This document contains screenshots from the CloudSecure AI platform, showcasing the various features and capabilities designed for small and medium businesses.\n").bold = True
    
    # Get all screenshot files
    screenshot_dir = "docs/screenshots"
    screenshot_files = glob.glob(os.path.join(screenshot_dir, "*.png"))
    
    # Define screenshot descriptions
    screenshot_descriptions = {
        "dashboard-overview.png": "Main Dashboard - Overview of cloud security posture with real-time metrics and alerts",
        "main-dashboard.png": "Dashboard Home - Central hub showing security score, recent alerts, and resource overview",
        "cloud-account-setup.png": "Cloud Account Setup - Connect your AWS, Azure, and GCP accounts in just a few clicks",
        "cloud-accounts.png": "Cloud Accounts Management - Manage all your cloud providers from one unified interface",
        "security-scan.png": "Security Scan Process - Watch as our AI scans your cloud infrastructure for vulnerabilities",
        "vulnerability-scanner.png": "Vulnerability Scanner - Real-time vulnerability detection with severity classification",
        "compliance-dashboard.png": "Compliance Dashboard - Track your compliance status across GDPR, SOC 2, and HIPAA frameworks",
        "compliance-reports.png": "Compliance Reports - Automated compliance reporting for audits and certifications",
        "ai-chat.png": "AI Chat Assistant - Get instant security advice from our AI-powered assistant",
        "ai-chat-interface.png": "AI Chat Interface - Interactive chat with AI for security guidance and recommendations",
        "system-settings.png": "System Settings - Configure your security preferences and notification settings"
    }
    
    # Process each screenshot
    for i, screenshot_path in enumerate(sorted(screenshot_files), 1):
        filename = os.path.basename(screenshot_path)
        
        # Add section heading
        heading = doc.add_heading(f'{i}. {filename.replace(".png", "").replace("-", " ").title()}', level=2)
        
        # Add description if available
        if filename in screenshot_descriptions:
            desc = doc.add_paragraph(screenshot_descriptions[filename])
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resize and add image
        try:
            resized_path = resize_image(screenshot_path)
            
            # Add the image to the document
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Add image with appropriate size
            doc.add_picture(resized_path, width=Inches(6), height=Inches(4))
            
            # Clean up resized image
            if resized_path != screenshot_path and os.path.exists(resized_path):
                os.remove(resized_path)
                
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            # Add placeholder text if image can't be processed
            error_para = doc.add_paragraph(f"[Screenshot: {filename}]")
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space between screenshots
        doc.add_paragraph()
    
    # Add footer information
    doc.add_page_break()
    footer_heading = doc.add_heading('About CloudSecure AI', level=1)
    footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("CloudSecure AI - Empowering Small Businesses with Enterprise-Grade Security\n\n").bold = True
    footer_text.add_run("A comprehensive, free cloud security platform designed specifically for small and medium businesses (SMBs) to protect their digital assets without the need for expensive IT expertise.\n\n")
    footer_text.add_run("Contact: duxinpi@gmail.com, ddxhvm@gmail.com")
    
    # Save the document
    doc.save('CloudSecure_AI_Screenshots.docx')
    print("Word document created: CloudSecure_AI_Screenshots.docx")

if __name__ == "__main__":
    try:
        create_word_document()
    except ImportError as e:
        print(f"Missing required library: {e}")
        print("Please install required packages:")
        print("pip install python-docx pillow")
    except Exception as e:
        print(f"Error creating Word document: {e}")
